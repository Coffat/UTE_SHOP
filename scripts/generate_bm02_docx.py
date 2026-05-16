#!/usr/bin/env python3
"""Generate BM02 Mentoring Agreement as DOCX — Times New Roman, formal layout."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


def set_run_font(run, name: str = "Times New Roman", size_pt: float = 12, bold: bool = False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    # Vietnamese / East Asian font hint for Word
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def set_cell_shading(cell, fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def apply_normal_style(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)


def add_heading_para(doc: Document, text: str, level: str = "section"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == "title" else 10)
    p.paragraph_format.space_after = Pt(6)
    if level == "title":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size_pt=14, bold=True)
    else:
        run = p.add_run(text)
        set_run_font(run, size_pt=12, bold=True)


def add_body(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        rest = text[len(bold_prefix) :]
        if rest:
            r2 = p.add_run(rest)
            set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)


def add_bullet(doc: Document, text: str, bold_label: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    if bold_label and text.startswith(bold_label):
        r1 = p.add_run(bold_label)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_label) :])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)


def main():
    root = Path(__file__).resolve().parents[1]
    out = root / "BM02-Mentoring-Agreement-Hoan-Thien.docx"

    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2)

    apply_normal_style(doc)

    # Header table
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for cell in t.rows[0].cells:
        cell.width = Cm(7.5)
        set_cell_shading(cell, "F2F2F2")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.clear()
    left_text = (
        "TRƯỜNG ĐẠI HỌC CÔNG NGHỆ KỸ THUẬT THÀNH PHỐ HỒ CHÍ MINH\n"
        "PHÒNG QUAN HỆ DOANH NGHIỆP"
    )
    right_text = (
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n"
        "Độc lập – Tự do – Hạnh phúc"
    )
    p0 = t.rows[0].cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in left_text.split("\n"):
        run = p0.add_run(line + ("\n" if line != left_text.split("\n")[-1] else ""))
        set_run_font(run, size_pt=11, bold=True)

    p1 = t.rows[0].cells[1].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in right_text.split("\n"):
        run = p1.add_run(line + ("\n" if line != right_text.split("\n")[-1] else ""))
        set_run_font(run, size_pt=11, bold=True)

    doc.add_paragraph()

    add_heading_para(doc, "THỎA THUẬN ĐỒNG HÀNH MENTOR & MENTEE", "title")

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("Mục đích: ")
    set_run_font(r, bold=True)
    r2 = p.add_run(
        "giúp mentor và mentee thống nhất kỳ vọng, cách làm việc và cam kết đồng hành."
    )
    set_run_font(r2)

    add_heading_para(doc, "A. Thông tin cặp mentor – mentee")

    bullets_a = [
        ("Họ và tên Mentor:", " Trần Quang Phúc"),
        ("Đơn vị/Công ty/Chức vụ:", " Công ty Ajinomoto Việt Nam"),
        ("Họ tên Mentee:", " Vũ Toàn Thắng"),
        ("Khoa/Ngành:", " Khoa Công nghệ Thông tin"),
        ("Thời gian bắt đầu mentoring:", " Tháng 5/2026"),
        ("Hình thức trao đổi:", " ☐ Trực tiếp    ☐ Online    ☑ Kết hợp"),
    ]
    for label, rest in bullets_a:
        add_bullet(doc, label + rest, bold_label=label)

    p_note = doc.add_paragraph()
    p_note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rn = p_note.add_run("(Đánh dấu ô phù hợp thực tế khi in và ký.)")
    set_run_font(rn, size_pt=11)
    rn.font.italic = True

    add_heading_para(doc, "B. Mục tiêu mentorship")
    goals = [
        (
            "Mục tiêu 1:",
            " Định hướng phát triển nghề nghiệp trong lĩnh vực CNTT và hiểu vận hành thực tế trong doanh nghiệp (quy trình, kỹ năng cần thiết, kỳ vọng nhà tuyển dụng).",
        ),
        (
            "Mục tiêu 2:",
            " Rèn kỹ năng mềm: giao tiếp chuyên nghiệp, làm việc nhóm, quản lý thời gian và trình bày ý tưởng.",
        ),
        (
            "Mục tiêu 3:",
            " Vận dụng kiến thức chuyên môn vào bài toán thực tế (ví dụ: thực tập/dự án), xây dựng kế hoạch học tập và hành động ngắn hạn rõ ràng sau mỗi buổi đồng hành.",
        ),
    ]
    for label, rest in goals:
        add_bullet(doc, label + rest, bold_label=label)

    add_heading_para(doc, "C. Cách thức làm việc")
    c_items = [
        (
            "Tần suất trao đổi:",
            " 01 buổi / 2 tuần (có thể điều chỉnh theo học kỳ và lịch hai bên); liên lạc ngắn qua kênh đã thống nhất khi cần làm rõ câu hỏi hoặc lịch hẹn.",
        ),
        (
            "Kênh liên lạc chính:",
            " Email và nền tảng họp trực tuyến Google Meet; tin nhắn (Zalo) cho việc hẹn giờ.",
        ),
        (
            "Thời lượng mỗi buổi:",
            " 45–60 phút (hoặc 90 phút nếu hai bên thống nhất cho buổi định kỳ).",
        ),
        (
            "Cách mentee chuẩn bị trước buổi gặp:",
            " Gửi agenda và 3–5 câu hỏi/chủ đề ưu tiên trước buổi ít nhất 24 giờ; tóm tắt tiến độ bài tập/dự án/thực tập và khó khăn gặp phải.",
        ),
        (
            "Cách mentor phản hồi / góp ý:",
            " Phản hồi trực tiếp trong buổi; tóm tắt 3–5 điểm chính và gợi ý hành động tiếp theo sau buổi (qua email ngắn nếu cần).",
        ),
        (
            "Cách ghi nhận nội dung sau buổi gặp:",
            " Mentee ghi biên bản ngắn (chủ đề đã trao đổi, cam kết hành động, thời hạn); mentor xác nhận hoặc bổ sung qua email/kênh chính trong vòng 3–5 ngày làm việc.",
        ),
    ]
    for label, rest in c_items:
        add_bullet(doc, label + rest, bold_label=label)

    add_heading_para(doc, "D. Cam kết của mentee")
    add_body(doc, "Mentee cam kết:")
    for line in [
        "Chủ động chuẩn bị trước mỗi buổi mentoring.",
        "Đúng giờ và tôn trọng thời gian của mentor.",
        "Chủ động đặt câu hỏi và chia sẻ khó khăn thật sự.",
        "Ghi nhận phản hồi với thái độ cầu thị.",
        "Thực hiện các hành động đã thống nhất.",
        "Hoàn thành reflection và biểu mẫu theo yêu cầu của chương trình.",
    ]:
        add_bullet(doc, line)

    add_heading_para(doc, "E. Cam kết của mentor")
    add_body(doc, "Mentor cam kết:")
    for line in [
        "Đồng hành, định hướng và góp ý trên tinh thần xây dựng.",
        "Tôn trọng cá tính, nền tảng và mục tiêu riêng của mentee.",
        "Chia sẻ kinh nghiệm thực tế phù hợp với nhu cầu của mentee.",
        "Khuyến khích mentee tự suy nghĩ, tự lựa chọn và tự chịu trách nhiệm với kế hoạch phát triển của mình.",
        "Phản hồi ngắn gọn về tiến độ của mentee theo biểu mẫu chương trình.",
    ]:
        add_bullet(doc, line)

    add_heading_para(doc, "F. Nguyên tắc chung")
    add_body(doc, "Hai bên thống nhất:")
    for line in [
        "Tôn trọng.",
        "Bảo mật những chia sẻ cá nhân khi cần thiết.",
        "Không phán xét.",
        "Đúng giờ.",
        "Phản hồi tích cực.",
        "Tập trung vào sự phát triển của mentee.",
    ]:
        add_bullet(doc, line)

    doc.add_paragraph()
    sig = doc.add_table(rows=3, cols=2)
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in sig.rows:
        for cell in row.cells:
            cell.width = Cm(7.5)

    def sig_cell(row_i: int, col_i: int, text: str, sub: str | None = None, bold_main: bool = True):
        cell = sig.rows[row_i].cells[col_i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        r = p.add_run(text)
        set_run_font(r, bold=bold_main)
        if sub:
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rs = p2.add_run(sub)
            set_run_font(rs, size_pt=11)
            rs.font.italic = True

    sig_cell(0, 0, "Mentor")
    sig_cell(0, 1, "Mentee")
    sig_cell(1, 0, "")
    sig_cell(1, 1, "")
    # Blank rows for signature
    sig.rows[1].cells[0].paragraphs[0].add_run("\n\n\n")
    sig.rows[1].cells[1].paragraphs[0].add_run("\n\n\n")
    sig_cell(2, 0, "(Ký tên, ghi họ tên)", bold_main=False)
    sig_cell(2, 1, "(Ký tên, ghi họ tên)", bold_main=False)

    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
